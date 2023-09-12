
import docx
import os
import pandas as pd
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx2pdf import convert

# Defining functions to add page number


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# Specifying file paths for excel, doc and pdf
workbook_path = 'D:\Archive.xlsx'
docx_path = 'D:\Archived Files.docx'
pdf_path = 'D:\Archived Files.pdf'

# Reading Archive excel for creating dataframe
df_xlsx = pd.read_excel(workbook_path)

topic = ''

# Creating docx object
d = docx.Document()

# Setting number of columns and column spacing
section = d.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '3')
cols.set(qn('w:space'), '12')

# Setting font style and font size
style = d.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(8)

# Setting margin widths
sections = d.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# Iterating over each row in dataframe
for i in range(0, len(df_xlsx)):

    # Creating subfolder titles based on Directory column value
    if pd.isnull(df_xlsx.at[i, 'Directory']) == False:
        topic_new = str(df_xlsx.at[i, 'Category']) + \
            '\\' + str(df_xlsx.at[i, 'Directory'])
    else:
        topic_new = str(df_xlsx.at[i, 'Category'])

    # Adding subfolder titles after styling
    if topic_new != topic:
        topic = topic_new
        p_topic = d.add_paragraph()
        p_topic.paragraph_format.space_before = Pt(8)
        p_topic.paragraph_format.space_after = Pt(1)
        run = p_topic.add_run(topic)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.underline = True

    # Assigning Content cell value to a variable for reference
    content = str(df_xlsx.at[i, 'Content'])

    # Highlighting unavailable Contents with Grey color
    if df_xlsx.at[i, 'Location'] == 'NT' or df_xlsx.at[i, 'Location'] == 'TBA':
        p_content = d.add_paragraph()
        run = p_content.add_run(content)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    # Adding Content without styling
    else:
        p_content = d.add_paragraph(content)

    # Specifying narrow spacing between Content lines
    p_content.paragraph_format.space_before = Pt(0.5)
    p_content.paragraph_format.space_after = Pt(0.5)

# Setting narrow spacing before first line
d.paragraphs[0].paragraph_format.space_before = Pt(0)

# Inserting line break
p_break = d.add_paragraph()
run = p_break.add_run()
run.add_break(WD_BREAK.LINE)

# Adding Date at the end of document
date = datetime.now()
p_date = d.add_paragraph()
run = p_date.add_run('ARCHIVED FILES DATED ' + date.strftime("%d %b %Y"))
run.bold = True

# Specifying alignment and spacing of page number in footer
d.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
d.sections[0].footer.paragraphs[0].paragraph_format.space_before = Pt(12)
add_page_number(d.sections[0].footer.paragraphs[0].add_run())

# Saving doc file
d.save(docx_path)

# Creating pdf file
convert(docx_path, pdf_path)

# Deleting doc file
os.remove(docx_path)
